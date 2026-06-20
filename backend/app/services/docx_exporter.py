from __future__ import annotations

import io
from typing import Any

# All python-docx imports are deferred inside generate_resume_docx so that
# this module can be imported in tests without the package installed.


def _set_heading_style(paragraph, font_size: int = 12, bold: bool = True) -> None:
    from docx.shared import Pt
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(font_size)


def _add_horizontal_rule(document) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _to_str(value: Any) -> str:
    """Safely convert a value to string, flattening dicts/lists if needed."""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        for key in ("text", "description", "name", "value", "content"):
            if key in value:
                return _to_str(value[key])
        return " ".join(_to_str(v) for v in value.values() if v)
    if isinstance(value, list):
        return " ".join(_to_str(v) for v in value if v)
    return str(value) if value is not None else ""


def _item_matches_label(item: dict, label: str, name_keys: tuple[str, ...]) -> bool:
    """Check whether a resume entry matches a rewrite's item_label.

    Used to target a specific entry (e.g. which project) when original_text
    is blank, since an empty string can't be matched by content.
    """
    if not label:
        return False
    label_lower = label.lower()
    item_text = " ".join(str(item.get(k, "")) for k in name_keys).lower()
    return bool(item_text) and (item_text in label_lower or label_lower in item_text)


def _apply_rewrites(resume_json: dict, approved_rewrites: list[dict]) -> dict:
    """Apply approved rewrite suggestions to the resume_json dict.

    Returns a copy of resume_json with suggested_text replacing original_text
    in the relevant section. When original_text is blank (a suggestion to
    fill in a missing description), item_label is used to target the
    specific entry instead of a content match.
    """
    import copy
    data = copy.deepcopy(resume_json)

    for rewrite in approved_rewrites:
        section = (rewrite.get("section") or "").lower()
        item_label = rewrite.get("item_label", "")
        original = rewrite.get("original_text", "")
        suggested = rewrite.get("suggested_text", "")

        if not suggested:
            continue

        if section in ("summary", "profile", "objective"):
            if not original:
                continue
            if isinstance(data.get("summary"), str):
                data["summary"] = data["summary"].replace(original, suggested)

        elif section in ("skills", "skill"):
            if not original:
                continue
            skills = data.get("skills", [])
            data["skills"] = [suggested if s == original else s for s in skills]

        elif section in ("work_experience", "experience", "work experience"):
            for exp in data.get("work_experience", []):
                if not original:
                    if _item_matches_label(exp, item_label, ("company", "title")):
                        exp["description"] = suggested
                    continue
                bullets = exp.get("bullets", [])
                exp["bullets"] = [suggested if b == original else b for b in bullets]
                if isinstance(exp.get("description"), str):
                    exp["description"] = exp["description"].replace(original, suggested)

        elif section in ("education",):
            if not original:
                continue
            for edu in data.get("education", []):
                if isinstance(edu.get("description"), str):
                    edu["description"] = edu["description"].replace(original, suggested)

        elif section in ("projects", "project"):
            for proj in data.get("projects", []):
                if not original:
                    if _item_matches_label(proj, item_label, ("name",)):
                        proj["description"] = suggested
                    continue
                bullets = proj.get("bullets", [])
                proj["bullets"] = [suggested if b == original else b for b in bullets]
                if isinstance(proj.get("description"), str):
                    proj["description"] = proj["description"].replace(original, suggested)

        elif section in ("leadership",):
            for item in data.get("leadership", []):
                if not original:
                    if _item_matches_label(item, item_label, ("title", "organization")):
                        item["description"] = suggested
                    continue
                bullets = item.get("bullets", [])
                item["bullets"] = [suggested if b == original else b for b in bullets]
                if isinstance(item.get("description"), str):
                    item["description"] = item["description"].replace(original, suggested)

        elif section in ("achievements", "achievement"):
            for ach in data.get("achievements", []):
                if not original:
                    if _item_matches_label(ach, item_label, ("title",)):
                        ach["description"] = suggested
                    continue
                if isinstance(ach.get("description"), str):
                    ach["description"] = ach["description"].replace(original, suggested)

        elif section in ("certifications", "certification"):
            for cert in data.get("certifications", []):
                if not original:
                    if _item_matches_label(cert, item_label, ("name",)):
                        cert["description"] = suggested
                    continue
                if isinstance(cert.get("description"), str):
                    cert["description"] = cert["description"].replace(original, suggested)

    return data


def generate_resume_docx(resume_json: dict[str, Any], approved_rewrites: list[dict]) -> bytes:
    """Generate a styled DOCX resume applying approved rewrite suggestions.

    Args:
        resume_json: Parsed resume data dict (matching ResumeJSON schema).
        approved_rewrites: List of approved rewrite suggestion dicts with keys:
            section, original_text, suggested_text, reason.

    Returns:
        DOCX file as bytes.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data = _apply_rewrites(resume_json, approved_rewrites)
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Header: Name & Contact ─────────────────────────────────────────────
    name = data.get("name", "")
    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(18)

    contact_parts: list[str] = []
    for field in ("email", "phone"):
        value = data.get(field, "")
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)

    # ── Summary ────────────────────────────────────────────────────────────
    summary = data.get("summary", "")
    if summary:
        doc.add_paragraph()
        heading = doc.add_paragraph("SUMMARY")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.space_after = Pt(4)

    # ── Skills ────────────────────────────────────────────────────────────
    skills: list = data.get("skills", [])
    if skills:
        doc.add_paragraph()
        heading = doc.add_paragraph("SKILLS")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        skills_para = doc.add_paragraph(", ".join(_to_str(s) for s in skills if s))
        skills_para.paragraph_format.space_after = Pt(4)

    # ── Education ─────────────────────────────────────────────────────────
    education: list[dict] = data.get("education", [])
    if education:
        doc.add_paragraph()
        heading = doc.add_paragraph("EDUCATION")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for edu in education:
            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            field = edu.get("field_of_study", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            gpa = edu.get("gpa")
            description = edu.get("description", "")

            date_str = f"{start} – {end}" if start or end else ""
            title_parts = [degree]
            if field:
                title_parts.append(f"in {field}")
            degree_str = " ".join(title_parts)

            # Institution line
            inst_para = doc.add_paragraph()
            inst_para.paragraph_format.space_before = Pt(4)
            inst_para.paragraph_format.space_after = Pt(0)
            inst_run = inst_para.add_run(institution)
            inst_run.bold = True
            if date_str:
                inst_para.add_run(f"  {date_str}").font.size = Pt(10)

            # Degree line
            if degree_str:
                deg_para = doc.add_paragraph(degree_str)
                deg_para.paragraph_format.space_before = Pt(0)
                deg_para.paragraph_format.space_after = Pt(0)

            if gpa:
                doc.add_paragraph(f"GPA: {gpa}").paragraph_format.space_after = Pt(0)
            if description:
                doc.add_paragraph(description).paragraph_format.space_after = Pt(0)

    # ── Work Experience ───────────────────────────────────────────────────
    work_experience: list[dict] = data.get("work_experience", [])
    if work_experience:
        doc.add_paragraph()
        heading = doc.add_paragraph("WORK EXPERIENCE")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for exp in work_experience:
            company = exp.get("company", "")
            title = exp.get("title", "")
            location_str = exp.get("location", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            is_current = exp.get("is_current", False)
            bullets: list[str] = exp.get("bullets", [])
            description = exp.get("description", "")

            end_display = "Present" if is_current else (end or "")
            date_str = f"{start} – {end_display}" if (start or end_display) else ""

            # Company + date line
            comp_para = doc.add_paragraph()
            comp_para.paragraph_format.space_before = Pt(6)
            comp_para.paragraph_format.space_after = Pt(0)
            comp_run = comp_para.add_run(company)
            comp_run.bold = True
            if date_str:
                comp_para.add_run(f"  {date_str}")

            # Title + location line
            if title or location_str:
                parts = [p for p in [title, location_str] if p]
                title_para = doc.add_paragraph("  |  ".join(parts))
                title_para.paragraph_format.space_before = Pt(0)
                title_para.paragraph_format.space_after = Pt(2)
                for run in title_para.runs:
                    run.italic = True

            # Bullet points
            for bullet in bullets:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.space_before = Pt(0)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.add_run(_to_str(bullet))

            # Description fallback
            if description and not bullets:
                desc_para = doc.add_paragraph(description)
                desc_para.paragraph_format.space_before = Pt(0)
                desc_para.paragraph_format.space_after = Pt(2)

    # ── Projects ──────────────────────────────────────────────────────────
    projects: list[dict] = data.get("projects", [])
    if projects:
        doc.add_paragraph()
        heading = doc.add_paragraph("PROJECTS")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for proj in projects:
            proj_name = proj.get("name", "")
            proj_desc = proj.get("description", "")
            technologies: list[str] = proj.get("technologies", [])
            url = proj.get("url", "")
            bullets_proj: list[str] = proj.get("bullets", [])
            start = proj.get("start_date", "")
            end = proj.get("end_date", "")
            date_str = f"{start} – {end}" if (start or end) else ""

            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(4)
            proj_para.paragraph_format.space_after = Pt(0)
            proj_run = proj_para.add_run(proj_name)
            proj_run.bold = True
            if date_str:
                proj_para.add_run(f"  {date_str}")

            if technologies:
                tech_para = doc.add_paragraph(f"Technologies: {', '.join(_to_str(t) for t in technologies if t)}")
                tech_para.paragraph_format.space_before = Pt(0)
                tech_para.paragraph_format.space_after = Pt(1)
                for run in tech_para.runs:
                    run.italic = True

            for bullet in bullets_proj:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.space_before = Pt(0)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.add_run(_to_str(bullet))

            if proj_desc and not bullets_proj:
                doc.add_paragraph(proj_desc)

            if url:
                url_para = doc.add_paragraph(f"URL: {url}")
                url_para.paragraph_format.space_after = Pt(2)

    # ── Leadership ────────────────────────────────────────────────────────
    leadership: list[dict] = data.get("leadership", [])
    if leadership:
        doc.add_paragraph()
        heading = doc.add_paragraph("LEADERSHIP")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for item in leadership:
            l_title = item.get("title", "")
            l_org = item.get("organization", "")
            l_start = item.get("start_date", "")
            l_end = item.get("end_date", "")
            l_desc = item.get("description", "")
            l_bullets: list[str] = item.get("bullets", [])
            date_str = f"{l_start} – {l_end}" if (l_start or l_end) else ""

            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(6)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run(l_title)
            title_run.bold = True
            if date_str:
                title_para.add_run(f"  {date_str}")

            if l_org:
                org_para = doc.add_paragraph(l_org)
                org_para.paragraph_format.space_before = Pt(0)
                org_para.paragraph_format.space_after = Pt(2)
                for run in org_para.runs:
                    run.italic = True

            for bullet in l_bullets:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.space_before = Pt(0)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.add_run(_to_str(bullet))

            if l_desc and not l_bullets:
                desc_para = doc.add_paragraph(l_desc)
                desc_para.paragraph_format.space_before = Pt(0)
                desc_para.paragraph_format.space_after = Pt(2)

    # ── Additional Sections ───────────────────────────────────────────────
    certifications: list = data.get("certifications", [])
    if certifications:
        doc.add_paragraph()
        heading = doc.add_paragraph("CERTIFICATIONS")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for cert in certifications:
            if isinstance(cert, dict):
                cert_name = cert.get("name", "")
                cert_issuer = cert.get("issuer", "")
                cert_date = cert.get("date", "")
                cert_desc = cert.get("description", "")

                cert_para = doc.add_paragraph()
                cert_para.paragraph_format.space_before = Pt(4)
                cert_para.paragraph_format.space_after = Pt(0)
                name_run = cert_para.add_run(cert_name)
                name_run.bold = True
                if cert_date:
                    cert_para.add_run(f"  {cert_date}")

                if cert_issuer:
                    issuer_para = doc.add_paragraph(cert_issuer)
                    issuer_para.paragraph_format.space_before = Pt(0)
                    issuer_para.paragraph_format.space_after = Pt(1)
                    for run in issuer_para.runs:
                        run.italic = True

                if cert_desc:
                    doc.add_paragraph(cert_desc).paragraph_format.space_after = Pt(2)
            else:
                doc.add_paragraph(_to_str(cert), style="List Bullet")

    achievements: list = data.get("achievements", [])
    if achievements:
        doc.add_paragraph()
        heading = doc.add_paragraph("ACHIEVEMENTS")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        for ach in achievements:
            if isinstance(ach, dict):
                ach_title = ach.get("title", "")
                ach_date = ach.get("date", "")
                ach_desc = ach.get("description", "")

                ach_para = doc.add_paragraph()
                ach_para.paragraph_format.space_before = Pt(4)
                ach_para.paragraph_format.space_after = Pt(0)
                title_run = ach_para.add_run(ach_title)
                title_run.bold = True
                if ach_date:
                    ach_para.add_run(f"  {ach_date}")

                if ach_desc:
                    doc.add_paragraph(ach_desc).paragraph_format.space_after = Pt(2)
            else:
                doc.add_paragraph(_to_str(ach), style="List Bullet")

    languages: list = data.get("languages", [])
    if languages:
        doc.add_paragraph()
        heading = doc.add_paragraph("LANGUAGES")
        _set_heading_style(heading, font_size=11, bold=True)
        _add_horizontal_rule(doc)
        lang_parts: list[str] = []
        for lang in languages:
            if isinstance(lang, dict):
                name = lang.get("language") or lang.get("name") or ""
                prof = lang.get("proficiency") or lang.get("level") or ""
                lang_parts.append(f"{name}: {prof}" if prof else name)
            else:
                lang_parts.append(_to_str(lang))
        doc.add_paragraph(", ".join(p for p in lang_parts if p))

    # ── Serialize to bytes ────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
