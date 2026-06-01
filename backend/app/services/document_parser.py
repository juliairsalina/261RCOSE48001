from __future__ import annotations

import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Dispatch to the correct extractor based on file_type.

    Args:
        file_bytes: Raw file bytes.
        file_type: MIME type or extension string, e.g. "application/pdf" or "pdf".

    Returns:
        Extracted plain text.

    Raises:
        ValueError: For unsupported file types.
    """
    normalized = file_type.lower().strip(".")
    if "pdf" in normalized:
        return extract_text_from_pdf(file_bytes)
    if "docx" in normalized or "openxmlformats" in normalized or "word" in normalized:
        return extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {file_type}. Only PDF and DOCX are supported.")
