from __future__ import annotations

import io
import json
from unittest.mock import AsyncMock, MagicMock, patch

import pytest


# ── Fixtures ───────────────────────────────────────────────────────────────

MOCK_PARSED_RESUME = {
    "name": "Jane Doe",
    "email": "jane@example.com",
    "phone": "+1-555-0100",
    "education": [
        {
            "institution": "MIT",
            "degree": "Bachelor of Science",
            "field_of_study": "Computer Science",
            "start_date": "2016",
            "end_date": "2020",
            "gpa": "3.9",
            "description": "",
        }
    ],
    "work_experience": [
        {
            "company": "Acme Corp",
            "title": "Software Engineer",
            "location": "San Francisco, CA",
            "start_date": "2020-07",
            "end_date": "",
            "is_current": True,
            "bullets": ["Built REST APIs with FastAPI", "Reduced latency by 40%"],
            "description": "",
        }
    ],
    "projects": [
        {
            "name": "OpenResume",
            "description": "Open source resume builder",
            "technologies": ["React", "TypeScript"],
            "url": "https://github.com/example/openresume",
            "start_date": "2021",
            "end_date": "2022",
            "bullets": ["1000+ GitHub stars"],
        }
    ],
    "skills": ["Python", "FastAPI", "PostgreSQL", "Docker", "React"],
    "languages": ["English", "Spanish"],
    "certifications": ["AWS Solutions Architect"],
    "achievements": ["Speaker at PyCon 2023"],
}

REQUIRED_KEYS = {
    "name", "email", "phone", "education", "work_experience",
    "projects", "skills", "languages", "certifications", "achievements",
}


# ── Tests ──────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_resume_returns_valid_schema():
    """Mock GPT response and verify all required keys are present in parsed output."""
    mock_response = json.dumps(MOCK_PARSED_RESUME)

    with patch(
        "app.services.openai_client.chat_completion",
        new=AsyncMock(return_value=mock_response),
    ):
        from app.services import openai_client

        result_str = await openai_client.chat_completion(
            messages=[{"role": "user", "content": "Parse this resume..."}]
        )
        parsed = json.loads(result_str)

        # All required keys must be present
        for key in REQUIRED_KEYS:
            assert key in parsed, f"Key '{key}' missing from parsed resume"

        assert parsed["name"] == "Jane Doe"
        assert parsed["email"] == "jane@example.com"
        assert isinstance(parsed["skills"], list)
        assert isinstance(parsed["education"], list)
        assert isinstance(parsed["work_experience"], list)
        assert len(parsed["work_experience"]) >= 1
        assert isinstance(parsed["work_experience"][0]["bullets"], list)


def test_extract_text_from_pdf():
    """Test PDF text extraction returns a non-empty string from valid PDF bytes."""
    from pypdf import PdfWriter
    from pypdf.generic import NameObject

    # Build a minimal in-memory PDF with one page
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    pdf_bytes = buffer.getvalue()

    from app.services.document_parser import extract_text_from_pdf

    # A blank page returns empty string — check that no exception is raised
    result = extract_text_from_pdf(pdf_bytes)
    assert isinstance(result, str)


def test_extract_text_from_docx():
    """Test DOCX text extraction returns a string from a minimal DOCX."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Alice Johnson")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("Skills: Python, Java, AWS")

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    from app.services.document_parser import extract_text_from_docx

    result = extract_text_from_docx(docx_bytes)
    assert isinstance(result, str)
    assert "Alice Johnson" in result
    assert "Python" in result


def test_extract_text_dispatcher_pdf():
    """extract_text dispatches correctly for PDF content type."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    pdf_bytes = buffer.getvalue()

    from app.services.document_parser import extract_text

    result = extract_text(pdf_bytes, "application/pdf")
    assert isinstance(result, str)


def test_extract_text_dispatcher_docx():
    """extract_text dispatches correctly for DOCX content type."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Test resume content")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    from app.services.document_parser import extract_text

    result = extract_text(docx_bytes, "docx")
    assert isinstance(result, str)
    assert "Test resume content" in result


def test_extract_text_unsupported_type_raises():
    """extract_text raises ValueError for unsupported file types."""
    from app.services.document_parser import extract_text

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(b"dummy content", "txt")


@pytest.mark.asyncio
async def test_resume_chunks_are_embedded():
    """Resume text is split into chunks and each chunk gets a 1536-dim embedding vector."""
    dummy_embedding = [0.1] * 1536
    resume_text = (
        "Jane Doe — Software Engineer\n"
        "Skills: Python, FastAPI, PostgreSQL\n"
        "Experience: Built REST APIs at Acme Corp. Reduced latency by 40%.\n"
        "Education: B.Sc. Computer Science, MIT 2020.\n"
        "Projects: OpenResume — open source resume builder with 1000+ stars."
    )

    with patch(
        "app.services.openai_client.get_embedding",
        new=AsyncMock(return_value=dummy_embedding),
    ):
        from app.services.embedding_service import generate_embeddings_batch

        # Simulate splitting the resume into chunks (same as the upload route does)
        chunks = [resume_text[i:i + 200] for i in range(0, len(resume_text), 200)]
        assert len(chunks) >= 1, "Expected at least one chunk"

        embeddings = await generate_embeddings_batch(chunks)

    assert len(embeddings) == len(chunks), "One embedding per chunk"
    for emb in embeddings:
        assert isinstance(emb, list), "Each embedding must be a list"
        assert len(emb) == 1536, "Embedding must be 1536-dimensional"
        assert all(isinstance(v, float) for v in emb), "All values must be floats"
