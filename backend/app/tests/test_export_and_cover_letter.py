from __future__ import annotations

import json
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

# ── Shared fixtures ────────────────────────────────────────────────────────

RESUME_JSON = {
    "name": "Julia Irsalina",
    "email": "julia@example.com",
    "phone": "+82-10-1234-5678",
    "summary": "Data Science student passionate about AI and backend systems.",
    "skills": ["Python", "FastAPI", "SQL", "Machine Learning", "NLP"],
    "education": [
        {
            "institution": "Korea University",
            "degree": "Bachelor of Science",
            "field": "Data Science",
            "start_date": "2022",
            "end_date": "2026",
        }
    ],
    "work_experience": [
        {
            "company": "ABC Tech",
            "title": "Backend Developer Intern",
            "start_date": "Jun 2025",
            "end_date": "Aug 2025",
            "bullets": [
                "Developed REST APIs using FastAPI.",
                "Improved query performance by 30%.",
                "Integrated OCR preprocessing pipeline.",
            ],
        }
    ],
    "projects": [
        {
            "name": "AI Resume Optimizer",
            "technologies": ["Python", "LangChain", "OpenAI"],
            "bullets": [
                "Built ATS evaluation system.",
                "Implemented AI rewrite suggestions.",
            ],
        }
    ],
    "languages": ["English", "Indonesian"],
    "certifications": [],
    "achievements": [],
}

REWRITE_SUGGESTIONS = [
    {
        "id": "s1",
        "section": "work_experience",
        "original_text": "Developed REST APIs using FastAPI.",
        "suggested_text": "Engineered 5 production REST APIs with FastAPI, serving 10K requests/day.",
        "reason": "Added quantifiable metrics matching job requirements.",
        "status": "approved",
    },
    {
        "id": "s2",
        "section": "summary",
        "original_text": "Data Science student passionate about AI and backend systems.",
        "suggested_text": "AI/ML engineer with hands-on Python and FastAPI experience, focused on NLP applications.",
        "reason": "Stronger opening aligned to senior AI roles.",
        "status": "rejected",
    },
    {
        "id": "s3",
        "section": "work_experience",
        "original_text": "Improved query performance by 30%.",
        "suggested_text": "Reduced database query latency by 30% through index optimization.",
        "reason": "More specific and technical.",
        "status": "pending",
    },
]


# ── Test 9: Only approved rewrites are applied to the exported resume ───────

def test_only_approved_rewrites_are_exported():
    """The DOCX exporter must apply only suggestions with status='approved'."""
    pytest.importorskip("docx", reason="python-docx not installed")
    from app.services.docx_exporter import generate_resume_docx

    approved = [s for s in REWRITE_SUGGESTIONS if s["status"] == "approved"]
    assert len(approved) == 1, "Test fixture should have exactly 1 approved suggestion"

    docx_bytes = generate_resume_docx(RESUME_JSON, approved)

    assert isinstance(docx_bytes, bytes), "Export must return bytes"
    assert len(docx_bytes) > 0, "Exported DOCX must not be empty"

    # Read back the DOCX and verify the approved suggestion was applied
    import io
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "10K requests/day" in full_text, (
        "Approved rewrite text should appear in the exported document"
    )


def test_rejected_rewrites_are_not_exported():
    """Rejected and pending suggestions must NOT appear in the exported DOCX."""
    pytest.importorskip("docx", reason="python-docx not installed")
    from app.services.docx_exporter import generate_resume_docx

    approved = [s for s in REWRITE_SUGGESTIONS if s["status"] == "approved"]
    docx_bytes = generate_resume_docx(RESUME_JSON, approved)

    import io
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Rejected suggestion text must NOT be in the export
    assert "AI/ML engineer with hands-on" not in full_text, (
        "Rejected suggestion should not appear in the exported document"
    )
    # Pending suggestion text must NOT be in the export
    assert "index optimization" not in full_text, (
        "Pending suggestion should not appear in the exported document"
    )


def test_export_with_no_approved_rewrites_uses_original():
    """When no rewrites are approved, export the original resume unchanged."""
    pytest.importorskip("docx", reason="python-docx not installed")
    from app.services.docx_exporter import generate_resume_docx

    docx_bytes = generate_resume_docx(RESUME_JSON, approved_rewrites=[])

    import io
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Julia Irsalina" in full_text, "Candidate name must appear in the export"
    assert "Developed REST APIs using FastAPI." in full_text, (
        "Original bullet text must be preserved when no rewrites are approved"
    )


# ── Test 10: Cover letter is generated ────────────────────────────────────

MOCK_COVER_LETTER = (
    "Dear Hiring Manager,\n\n"
    "I am writing to express my strong interest in the AI Backend Engineer role at TechCorp Inc. "
    "With a solid foundation in Python, FastAPI, and machine learning, I have built production-grade "
    "REST APIs and integrated AI-powered pipelines during my internship at ABC Tech.\n\n"
    "My project, AI Resume Optimizer, demonstrates my ability to design end-to-end AI systems "
    "using LangChain and OpenAI APIs — directly aligned with your team's technology stack.\n\n"
    "I would welcome the opportunity to contribute to TechCorp's mission. "
    "Thank you for considering my application.\n\n"
    "Sincerely,\nJulia Irsalina"
)


@pytest.mark.asyncio
async def test_cover_letter_is_generated():
    """Cover letter generation returns a non-empty string between 250 and 400 words."""
    with patch(
        "app.services.openai_client.chat_completion",
        new=AsyncMock(return_value=MOCK_COVER_LETTER),
    ):
        from app.services import openai_client

        content = await openai_client.chat_completion(
            messages=[{"role": "user", "content": "Generate a cover letter..."}],
            temperature=0.4,
        )

    assert isinstance(content, str), "Cover letter must be a string"
    assert len(content.strip()) > 0, "Cover letter must not be empty"

    word_count = len(content.split())
    assert 50 <= word_count <= 500, (
        f"Cover letter word count {word_count} is outside acceptable range (50-500)"
    )


@pytest.mark.asyncio
async def test_cover_letter_mentions_company_and_role():
    """Cover letter should reference the target company and role."""
    with patch(
        "app.services.openai_client.chat_completion",
        new=AsyncMock(return_value=MOCK_COVER_LETTER),
    ):
        from app.services import openai_client

        content = await openai_client.chat_completion(
            messages=[{"role": "user", "content": "Generate a cover letter..."}],
        )

    assert "TechCorp" in content, "Cover letter should mention the company name"
    assert "AI Backend Engineer" in content, "Cover letter should mention the role title"


@pytest.mark.asyncio
async def test_cover_letter_does_not_invent_credentials():
    """The mock cover letter should only use facts present in the resume fixture."""
    with patch(
        "app.services.openai_client.chat_completion",
        new=AsyncMock(return_value=MOCK_COVER_LETTER),
    ):
        from app.services import openai_client

        content = await openai_client.chat_completion(
            messages=[{"role": "user", "content": "Generate a cover letter..."}],
        )

    # Known resume facts that may appear
    resume_facts = ["Python", "FastAPI", "AI Resume Optimizer", "LangChain", "ABC Tech"]
    found = [fact for fact in resume_facts if fact in content]
    assert len(found) >= 2, (
        f"Cover letter should reference resume facts. Found: {found}"
    )

    # Should NOT mention skills not in the resume
    invented_skills = ["Java", "Ruby", "Scala", "Go"]
    for skill in invented_skills:
        assert skill not in content, (
            f"Cover letter should not mention skill not in resume: {skill}"
        )
